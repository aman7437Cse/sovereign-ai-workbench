import os
import pypdf
import openpyxl
from PIL import Image, ImageDraw, ImageFont
from docx import Document

def generate_sample_files(demo_dir: str):
    os.makedirs(demo_dir, exist_ok=True)

    # 1. Sample Inspection Report PDF
    pdf_path = os.path.join(demo_dir, "sample_inspection_report.pdf")
    doc = Document()
    doc.add_heading("NOVA INDUSTRIAL SYSTEMS - REFINERY UNIT 07", level=1)
    doc.add_heading("Ultrasonic & Mechanical Inspection Report", level=2)
    doc.add_paragraph("Date: 2025-08-20 | Inspector: J. Miller, Lead Mechanical Engineer")
    doc.add_paragraph("Equipment Identifier: Pipeline Loop 4B / Secondary Pressure Valve V-102")
    doc.add_paragraph(
        "Observation Summary:\n"
        "- Ultrasonic wall thickness measurement recorded at 8.76 mm (Nominal Baseline: 10.0 mm).\n"
        "- Localized wall thinning calculated at 12.4%.\n"
        "- Secondary Pressure Valve V-102 operating pressure measured at 41.8 bar (Max Allowable Limit: 45.0 bar).\n"
        "- Micro-visual check shows minor surface oxidation on lower flange body; zero active fluid leakage detected."
    )
    temp_docx = os.path.join(demo_dir, "_temp_report.docx")
    doc.save(temp_docx)

    # Convert simple text to PDF using pypdf / plain write or simple pdf builder
    writer = pypdf.PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    # Write PDF stream or text
    with open(pdf_path, "wb") as f:
        writer.write(f)
    
    # Save a text-rich PDF content
    with open(os.path.join(demo_dir, "sample_inspection_report.txt"), "w", encoding="utf-8") as f:
        f.write(
            "NOVA INDUSTRIAL SYSTEMS - REFINERY UNIT 07\n"
            "Ultrasonic & Mechanical Inspection Report\n"
            "Date: 2025-08-20 | Inspector: J. Miller, Lead Mechanical Engineer\n"
            "Equipment Identifier: Pipeline Loop 4B / Secondary Pressure Valve V-102\n\n"
            "Observation Summary:\n"
            "- Ultrasonic wall thickness measurement recorded at 8.76 mm (Nominal Baseline: 10.0 mm).\n"
            "- Localized wall thinning calculated at 12.4%.\n"
            "- Secondary Pressure Valve V-102 operating pressure measured at 41.8 bar (Max Allowable Limit: 45.0 bar).\n"
            "- Micro-visual check shows minor surface oxidation on lower flange body; zero active fluid leakage detected.\n"
        )

    # 2. Sample SOP Document
    sop_path = os.path.join(demo_dir, "sample_sop.pdf")
    with open(os.path.join(demo_dir, "sample_sop.txt"), "w", encoding="utf-8") as f:
        f.write(
            "STANDARD OPERATING PROCEDURE: #SOP-2025-07\n"
            "TITLE: REFINERY PIPELINE INSPECTION & OPERATIONAL CLEARANCE POLICY\n"
            "ORGANIZATION: NOVA INDUSTRIAL SYSTEMS\n\n"
            "SECTION 4.2: WALL THINNING & PRESSURE BOUND TOLERANCES\n"
            "1. Wall thinning <= 10.0%: Unconditional operational clearance.\n"
            "2. Wall thinning > 10.0% and <= 15.0%: Conditional operational clearance permitted subject to "
            "scheduled seal replacement within 30 days and monthly ultrasonic inspection logs.\n"
            "3. Wall thinning > 15.0%: Immediate shutdown required for emergency line replacement.\n\n"
            "SECTION 5.1: APPROVAL REQUIREMENTS\n"
            "All conditional clearance notes must be verified against ultrasonic logs and signed by the Chief Mechanical Engineer.\n"
        )

    # 3. Sample Engineering Drawing Image (PNG)
    img_path = os.path.join(demo_dir, "sample_engineering_image.png")
    img = Image.new("RGB", (800, 600), color=(15, 23, 42))
    draw = ImageDraw.Draw(img)

    # Draw industrial valve schematic shapes
    draw.rectangle([100, 100, 700, 500], outline=(56, 189, 248), width=3)
    draw.ellipse([300, 200, 500, 400], outline=(234, 179, 8), width=4)
    draw.rectangle([380, 150, 420, 450], fill=(71, 85, 105), outline=(255, 255, 255), width=2)
    draw.text((120, 120), "REFINERY UNIT-07: VALVE V-102 DIAGRAM", fill=(255, 255, 255))
    draw.text((320, 280), "VALVE V-102\n45 BAR LIMIT", fill=(255, 255, 255))
    draw.text((520, 420), "[CORROSION ZONE: 12.4%]", fill=(239, 68, 68))
    img.save(img_path)

    # 4. Sample Spreadsheet XLSX
    xlsx_path = os.path.join(demo_dir, "sample_data.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Maintenance Budget"
    ws.append(["Line Item ID", "Asset Description", "Plant Section", "Estimated Cost ($)", "Priority"])
    ws.append(["M-101", "Pressure Safety Valve V-102 Seal", "Unit 07 Refinery", 12500, "High"])
    ws.append(["M-102", "Pipeline Loop 4B Ultrasonic Check", "Unit 07 Refinery", 28400, "High"])
    ws.append(["M-103", "Cooling Pump P-104 Bearing Clean", "Unit 07 Refinery", 8900, "Medium"])
    wb.save(xlsx_path)

    # 5. Sample Coding Task TXT
    code_task_path = os.path.join(demo_dir, "sample_coding_task.txt")
    with open(code_task_path, "w", encoding="utf-8") as f:
        f.write("Write a Python script to validate telemetry pressure data and check against 45 bar threshold.")

    if os.path.exists(temp_docx):
        os.remove(temp_docx)

if __name__ == "__main__":
    generate_sample_files(os.path.abspath(os.path.join(os.path.dirname(__file__), "sample_files")))
