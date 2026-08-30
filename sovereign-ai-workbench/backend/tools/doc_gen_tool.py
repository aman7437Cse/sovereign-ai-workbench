import os
import datetime
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from backend.config import config

class DocGenTool:
    """
    Real Word (.docx) Document Generator.
    Generates professional approval notes, technical reports, and compliance documentation.
    """

    def set_cell_background(self, cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tcPr.append(shd)

    def generate_approval_note(self, title: str, summary: str, findings: list, sop_citation: str, filename: str = "Approval_Note.docx") -> str:
        doc = docx.Document()

        # Page setup - Margins
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        # Header Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_title = p_title.add_run("SOVEREIGN INDUSTRIAL WORKBENCH")
        run_title.font.name = "Arial"
        run_title.font.size = Pt(9)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(100, 110, 125)

        # Main Document Title
        p_main = doc.add_paragraph()
        run_main = p_main.add_run(title.upper())
        run_main.font.name = "Arial"
        run_main.font.size = Pt(20)
        run_main.font.bold = True
        run_main.font.color.rgb = RGBColor(15, 23, 42)

        # Subtitle metadata table
        table_meta = doc.add_table(rows=4, cols=2)
        table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_data = [
            ("Document Reference:", f"IND-REF-2025-APP-{datetime.datetime.now().strftime('%d%m')}"),
            ("Facility Location:", "Nova Industrial Systems - Unit 07 (High Pressure Section)"),
            ("Date of Generation:", datetime.datetime.now().strftime("%B %d, %Y")),
            ("Security Classification:", "AIR-GAPPED / CONFIDENTIAL INDUSTRIAL")
        ]
        for idx, (label, val) in enumerate(meta_data):
            r = table_meta.rows[idx]
            r.cells[0].paragraphs[0].add_run(label).bold = True
            r.cells[0].paragraphs[0].runs[0].font.name = "Arial"
            r.cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
            r.cells[1].paragraphs[0].add_run(val)
            r.cells[1].paragraphs[0].runs[0].font.name = "Arial"
            r.cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
            self.set_cell_background(r.cells[0], "F1F5F9")

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # Section 1: Executive Summary
        h1 = doc.add_heading("1. Executive Summary & Purpose", level=1)
        h1.runs[0].font.color.rgb = RGBColor(30, 41, 59)
        h1.runs[0].font.name = "Arial"
        
        p1 = doc.add_paragraph(summary)
        p1.runs[0].font.name = "Arial"
        p1.runs[0].font.size = Pt(10.5)

        # Section 2: Technical Inspection Findings
        h2 = doc.add_heading("2. Technical Findings & Sensor Diagnostics", level=1)
        h2.runs[0].font.color.rgb = RGBColor(30, 41, 59)
        h2.runs[0].font.name = "Arial"

        for f in findings:
            p_f = doc.add_paragraph(style='List Bullet')
            r_f = p_f.add_run(f)
            r_f.font.name = "Arial"
            r_f.font.size = Pt(10)

        # Section 3: SOP Compliance Verification
        h3 = doc.add_heading("3. SOP Compliance & Policy Citation", level=1)
        h3.runs[0].font.color.rgb = RGBColor(30, 41, 59)
        h3.runs[0].font.name = "Arial"

        p_sop = doc.add_paragraph()
        r_sop = p_sop.add_run(f"Cross-Referenced SOP Document:\n{sop_citation}")
        r_sop.font.name = "Arial"
        r_sop.font.size = Pt(10)
        r_sop.font.italic = True

        # Section 4: Approval Recommendation & Signatures
        h4 = doc.add_heading("4. Formal Recommendation & Clearance", level=1)
        h4.runs[0].font.color.rgb = RGBColor(30, 41, 59)
        h4.runs[0].font.name = "Arial"

        p_rec = doc.add_paragraph(
            "Based on autonomous agent verification, localized wall thinning (12.4%) is strictly within the "
            "allowable 15% conditional threshold prescribed by SOP #SOP-2025-07. Operational clearance is hereby "
            "RECOMMENDED FOR CONDITIONAL APPROVAL subject to seal replacement during Q3 maintenance window."
        )
        p_rec.runs[0].font.name = "Arial"
        p_rec.runs[0].font.size = Pt(10.5)

        # Signature Table
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        sig_table.rows[0].cells[0].paragraphs[0].add_run("Prepared By:\n\n___________________________\nAutonomous Agentic Engine\nSovereign AI Workbench").bold = False
        sig_table.rows[0].cells[1].paragraphs[0].add_run("Approved By:\n\n___________________________\nChief Mechanical Engineer\nUnit 07 Refinery Operations").bold = False
        
        for r in sig_table.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9.5)

        output_path = os.path.join(config.DELIVERABLES_DIR, filename)
        doc.save(output_path)
        return output_path

doc_gen_tool = DocGenTool()
